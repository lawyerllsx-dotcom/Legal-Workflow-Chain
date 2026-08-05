#!/usr/bin/env python3
"""Generate a generic evidence catalog DOCX from user-provided item data.

This public script intentionally contains no real case facts, parties, docket
numbers, property details, amounts, addresses, or local paths. Missing user
data is represented with explicit placeholders instead of guessed values.
"""

from __future__ import annotations

import argparse
import csv
import json
import sys
from pathlib import Path
from typing import Any

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
except ImportError as exc:  # pragma: no cover - depends on runtime setup
    raise SystemExit(
        "Missing dependency: python-docx. Install it with `python3 -m pip install python-docx`."
    ) from exc


FIELD_ALIASES = {
    "number": ["编号", "序号", "number", "no", "index"],
    "name": ["证据名称", "名称", "name", "title", "evidence_name"],
    "source": ["证据来源", "来源", "source", "provider"],
    "purpose": ["证明事项", "证明目的", "purpose", "fact_to_prove", "description"],
    "pages": ["页码", "页数", "页码范围", "pages", "page", "page_range"],
}

PLACEHOLDERS = {
    "name": "[待补: 证据名称]",
    "source": "[待补: 证据来源]",
    "purpose": "[待补: 证明事项]",
    "pages": "[待补: 页码]",
}


def _clean(value: Any) -> str:
    if value is None:
        return ""
    return str(value).strip()


def _lookup(row: dict[str, Any], canonical: str) -> str:
    lowered = {str(key).strip().lower(): value for key, value in row.items()}
    for alias in FIELD_ALIASES[canonical]:
        if alias in row:
            return _clean(row[alias])
        value = lowered.get(alias.lower())
        if value is not None:
            return _clean(value)
    return ""


def normalize_item(row: dict[str, Any], index: int, start_number: int) -> dict[str, str]:
    number = _lookup(row, "number") or str(start_number + index)
    return {
        "number": number,
        "name": _lookup(row, "name") or PLACEHOLDERS["name"],
        "source": _lookup(row, "source") or PLACEHOLDERS["source"],
        "purpose": _lookup(row, "purpose") or PLACEHOLDERS["purpose"],
        "pages": _lookup(row, "pages") or PLACEHOLDERS["pages"],
    }


def read_items(input_path: str | Path, start_number: int = 1) -> list[dict[str, str]]:
    path = Path(input_path)
    suffix = path.suffix.lower()

    if suffix == ".json":
        raw = json.loads(path.read_text(encoding="utf-8"))
        if not isinstance(raw, list):
            raise ValueError("JSON input must be an array of evidence item objects.")
        rows = raw
    elif suffix == ".csv":
        with path.open("r", encoding="utf-8-sig", newline="") as handle:
            rows = list(csv.DictReader(handle))
    elif suffix == ".xlsx":
        try:
            from openpyxl import load_workbook
        except ImportError as exc:  # pragma: no cover - optional dependency
            raise SystemExit(
                "XLSX input requires openpyxl. Install it with `python3 -m pip install openpyxl`."
            ) from exc
        workbook = load_workbook(path, data_only=True)
        sheet = workbook.active
        values = list(sheet.iter_rows(values_only=True))
        if not values:
            rows = []
        else:
            headers = [_clean(cell) for cell in values[0]]
            rows = [dict(zip(headers, row)) for row in values[1:] if any(_clean(cell) for cell in row)]
    else:
        raise ValueError("Unsupported input format. Use .csv, .json, or .xlsx.")

    return [normalize_item(dict(row), index, start_number) for index, row in enumerate(rows)]


def _replace_in_paragraph(paragraph, replacements: dict[str, str]) -> None:
    text = paragraph.text
    changed = False
    for old, new in replacements.items():
        if old in text:
            text = text.replace(old, new)
            changed = True
    if not changed:
        return
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def _replace_placeholders(document: Document, replacements: dict[str, str]) -> None:
    for paragraph in document.paragraphs:
        _replace_in_paragraph(paragraph, replacements)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, replacements)


def _find_catalog_table(document: Document):
    for table in document.tables:
        if not table.rows:
            continue
        header = "\n".join(cell.text for cell in table.rows[0].cells)
        if "编号" in header and "证据名称" in header:
            return table
    return document.tables[0] if document.tables else None


def _clear_body_rows(table, header_rows: int = 1) -> None:
    for row in list(table.rows)[header_rows:]:
        table._tbl.remove(row._tr)


def _ensure_table(document: Document):
    table = _find_catalog_table(document)
    if table is None:
        table = document.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        headers = ["编号", "证据名称", "证据来源", "证明事项", "页码"]
        for cell, header in zip(table.rows[0].cells, headers):
            cell.text = header
    return table


def _fill_table(table, items: list[dict[str, str]]) -> None:
    _clear_body_rows(table)
    headers = ["number", "name", "source", "purpose", "pages"]
    for item in items:
        cells = table.add_row().cells
        for cell, key in zip(cells, headers):
            cell.text = item[key]


def _new_document(
    title: str,
    case_name: str,
    docket_number: str,
    submitter: str,
    submit_date: str,
) -> Document:
    document = Document()
    heading = document.add_heading(title or "证据目录", level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for label, value in [
        ("案件名称", case_name or "[案件名称]"),
        ("案号", docket_number or "[案号]"),
        ("提交主体", submitter or "[提交主体]"),
        ("提交日期", submit_date or "[提交日期]"),
    ]:
        paragraph = document.add_paragraph()
        paragraph.add_run(f"{label}: ").bold = True
        paragraph.add_run(value)

    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for cell, header in zip(table.rows[0].cells, ["编号", "证据名称", "证据来源", "证明事项", "页码"]):
        cell.text = header
    return document


def _set_metadata(document: Document) -> None:
    props = document.core_properties
    props.author = "Legal AI Skills Contributors"
    props.last_modified_by = "Legal AI Skills Contributors"
    props.title = "Evidence Catalog Template"
    props.subject = "Evidence catalog generated from user-provided data"
    props.comments = "No real case facts are bundled with this public skill."


def generate_evidence_catalog(
    items: list[dict[str, str]],
    output_path: str | Path,
    template_path: str | Path | None = None,
    title: str = "证据目录",
    case_name: str = "[案件名称]",
    docket_number: str = "[案号]",
    submitter: str = "[提交主体]",
    submit_date: str = "[提交日期]",
) -> Path:
    if template_path:
        document = Document(str(template_path))
        _replace_placeholders(
            document,
            {
                "[标题]": title,
                "[案件名称]": case_name,
                "[案号]": docket_number,
                "[提交主体]": submitter,
                "[提交日期]": submit_date,
            },
        )
    else:
        document = _new_document(title, case_name, docket_number, submitter, submit_date)

    table = _ensure_table(document)
    _fill_table(table, items)

    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run.font.size = run.font.size or Pt(12)

    _set_metadata(document)
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    return output


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Generate a generic evidence catalog DOCX.")
    parser.add_argument("--input", required=True, help="CSV, JSON, or XLSX evidence item data.")
    parser.add_argument("--output", required=True, help="Output DOCX path.")
    parser.add_argument("--template", help="Optional user-provided DOCX template.")
    parser.add_argument("--title", default="证据目录")
    parser.add_argument("--case-name", default="[案件名称]")
    parser.add_argument("--docket-number", default="[案号]")
    parser.add_argument("--submitter", default="[提交主体]")
    parser.add_argument("--submit-date", default="[提交日期]")
    parser.add_argument("--start-number", type=int, default=1)
    return parser


def main(argv: list[str] | None = None) -> int:
    args = build_parser().parse_args(argv)
    items = read_items(args.input, start_number=args.start_number)
    if not items:
        raise SystemExit("No evidence items found in input data.")
    output = generate_evidence_catalog(
        items=items,
        output_path=args.output,
        template_path=args.template,
        title=args.title,
        case_name=args.case_name,
        docket_number=args.docket_number,
        submitter=args.submitter,
        submit_date=args.submit_date,
    )
    print(f"Generated evidence catalog: {output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
